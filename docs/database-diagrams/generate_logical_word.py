"""
Enlarge and enhance the logical data model for Microsoft Word insertion.

Outputs:
  02-logical-overview-word.png   — high-res 300 DPI PNG
  02-logical-overview-word.svg   — vector original
  02-logical-overview-word.docx  — landscape Word doc with the figure embedded
  Also refreshes 02-logical-overview.png/svg as the enlarged master.
"""

from __future__ import annotations

import math
import os
from dataclasses import dataclass, field
from html import escape
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm, RGBColor

OUT = Path(__file__).resolve().parent

# Print-friendly palette
BG = "#FFFFFF"
INK = "#1A1A1A"
MUTED = "#555555"
LINE = "#333333"
STROKE = "#B0B0B0"
WHITE = "#FFFFFF"
PK_TINT = "#FFF4D6"
FK_TINT = "#E8F1FA"
BAND = "#F3F4F6"

DOMAINS = {
    "identity": "#1B365D",
    "hr": "#0E4D4A",
    "payroll": "#6B3F1D",
    "academics": "#1B4D3E",
    "messaging": "#3D2B56",
    "governance": "#374151",
}

# Enlarged metrics for Word readability
HEADER_H = 42
ROW_H = 26
COL_W = 300
FONT = "Calibri, Segoe UI, Arial, sans-serif"
DPI = 300


def _font_path(bold: bool = False) -> str | None:
    windir = os.environ.get("WINDIR", r"C:\Windows")
    names = (
        ["calibrib.ttf", "segoeuib.ttf", "arialbd.ttf"]
        if bold
        else ["calibri.ttf", "segoeui.ttf", "arial.ttf", "tahoma.ttf"]
    )
    for name in names:
        path = Path(windir) / "Fonts" / name
        if path.exists():
            return str(path)
    return None


def _pil_font(size: int, bold: bool = False):
    path = _font_path(bold)
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


@dataclass
class Col:
    name: str
    typ: str
    pk: bool = False
    fk: bool = False
    uk: bool = False


@dataclass
class Table:
    name: str
    domain: str
    cols: list[Col]
    x: float
    y: float
    width: float = COL_W

    @property
    def h(self) -> float:
        return HEADER_H + ROW_H * max(len(self.cols), 1)

    @property
    def right(self) -> float:
        return self.x + self.width

    @property
    def bottom(self) -> float:
        return self.y + self.h

    @property
    def cx(self) -> float:
        return self.x + self.width / 2

    @property
    def cy(self) -> float:
        return self.y + self.h / 2

    def col_y(self, name: str) -> float:
        for i, col in enumerate(self.cols):
            if col.name == name:
                return self.y + HEADER_H + i * ROW_H + ROW_H / 2
        return self.cy


@dataclass
class Rel:
    a: str
    a_col: str
    b: str
    b_col: str
    a_card: str
    b_card: str
    a_side: str = "right"
    b_side: str = "left"
    label: str = ""
    via: list[tuple[float, float]] = field(default_factory=list)


class Canvas:
    def __init__(self, width: int, height: int, title: str, subtitle: str):
        self.width = width
        self.height = height
        self.title = title
        self.subtitle = subtitle
        self.svg: list[str] = []
        self.img = Image.new("RGB", (width, height), BG)
        self.draw = ImageDraw.Draw(self.img)
        self.f12 = _pil_font(12)
        self.f13 = _pil_font(13)
        self.f14 = _pil_font(14)
        self.f14b = _pil_font(14, True)
        self.f16 = _pil_font(16)
        self.f16b = _pil_font(16, True)
        self.f18b = _pil_font(18, True)
        self.f22b = _pil_font(22, True)
        self.f28b = _pil_font(28, True)

    def rect(self, x, y, w, h, fill, stroke=STROKE, sw=1, radius=0):
        svg_fill = fill if fill and fill != "none" else "none"
        self.svg.append(
            f'<rect x="{x:.1f}" y="{y:.1f}" width="{w:.1f}" height="{h:.1f}" '
            f'rx="{radius}" fill="{svg_fill}" stroke="{stroke}" stroke-width="{sw}"/>'
        )
        pil_fill = None if svg_fill == "none" else fill
        box = [x, y, x + w, y + h]
        if radius:
            self.draw.rounded_rectangle(box, radius=radius, fill=pil_fill, outline=stroke, width=max(1, round(sw)))
        else:
            self.draw.rectangle(box, fill=pil_fill, outline=stroke, width=max(1, round(sw)))

    def line(self, pts, color=LINE, sw=1.8, dashed=False):
        d = " ".join(f"{'M' if i == 0 else 'L'}{x:.1f},{y:.1f}" for i, (x, y) in enumerate(pts))
        dash = ' stroke-dasharray="6 5"' if dashed else ""
        self.svg.append(
            f'<path d="{d}" fill="none" stroke="{color}" stroke-width="{sw}" '
            f'stroke-linejoin="round" stroke-linecap="round"{dash}/>'
        )
        xy = [(round(x), round(y)) for x, y in pts]
        if dashed:
            self._dashed(xy, color, sw)
        else:
            self.draw.line(xy, fill=color, width=max(1, round(sw)))

    def _dashed(self, xy, color, sw):
        on, rem = True, 0.0
        for i in range(len(xy) - 1):
            x1, y1 = xy[i]
            x2, y2 = xy[i + 1]
            dx, dy = x2 - x1, y2 - y1
            dist = math.hypot(dx, dy) or 1
            ux, uy = dx / dist, dy / dist
            pos = 0.0
            cx, cy = float(x1), float(y1)
            while pos < dist:
                seg = (6.0 if on else 5.0) - rem
                take = min(seg, dist - pos)
                nx, ny = cx + ux * take, cy + uy * take
                if on:
                    self.draw.line([(cx, cy), (nx, ny)], fill=color, width=max(1, round(sw)))
                rem = 0
                if take < seg:
                    rem = seg - take
                    cx, cy = nx, ny
                    break
                on = not on
                cx, cy = nx, ny
                pos += take

    def circle(self, x, y, r, fill=WHITE, stroke=LINE, sw=1.6):
        self.svg.append(
            f'<circle cx="{x:.1f}" cy="{y:.1f}" r="{r:.1f}" fill="{fill}" '
            f'stroke="{stroke}" stroke-width="{sw}"/>'
        )
        self.draw.ellipse([x - r, y - r, x + r, y + r], fill=fill, outline=stroke, width=max(1, round(sw)))

    def text(self, x, y, s, size=14, color=INK, bold=False, anchor="start"):
        weight = "600" if bold else "400"
        anc = {"start": "start", "middle": "middle", "end": "end"}[anchor]
        self.svg.append(
            f'<text x="{x:.1f}" y="{y:.1f}" fill="{color}" font-size="{size}" '
            f'font-family="{FONT}" font-weight="{weight}" text-anchor="{anc}" '
            f'dominant-baseline="central">{escape(s)}</text>'
        )
        if size >= 26:
            font = self.f28b
        elif size >= 20:
            font = self.f22b
        elif size >= 17:
            font = self.f18b
        elif bold:
            font = self.f16b if size >= 15 else self.f14b
        else:
            font = self.f16 if size >= 15 else (self.f14 if size >= 14 else (self.f13 if size >= 13 else self.f12))
        ax = {"start": "lm", "middle": "mm", "end": "rm"}[anchor]
        self.draw.text((x, y), s, fill=color, font=font, anchor=ax)

    def save(self, stem: str):
        header = (
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{self.width}" height="{self.height}" '
            f'viewBox="0 0 {self.width} {self.height}">\n'
            f'<rect width="100%" height="100%" fill="{BG}"/>\n'
        )
        svg_path = OUT / f"{stem}.svg"
        svg_path.write_text(header + "\n".join(self.svg) + "\n</svg>\n", encoding="utf-8")
        png_path = OUT / f"{stem}.png"
        self.img.save(png_path, "PNG", dpi=(DPI, DPI))
        return svg_path, png_path


PK = lambda n="id", t="bigint": Col(n, t, pk=True)
FK = lambda n, t="bigint": Col(n, t, fk=True)
UK = lambda n, t="varchar": Col(n, t, uk=True)
C = lambda n, t="varchar": Col(n, t)


def draw_title(c: Canvas):
    c.text(40, 36, c.title, size=28, color=INK, bold=True)
    c.text(40, 68, c.subtitle, size=15, color=MUTED)
    c.line([(40, 88), (c.width - 40, 88)], color="#D0D0D0", sw=1.5)


def draw_domain_banner(c: Canvas, x, y, w, title, domain):
    color = DOMAINS[domain]
    c.rect(x, y, w, 34, color, color, 0, radius=6)
    c.text(x + w / 2, y + 17, title, size=14, color=WHITE, bold=True, anchor="middle")


def draw_legend(c: Canvas, x: float, y: float):
    c.rect(x, y, 900, 72, WHITE, STROKE, 1.2, radius=8)
    c.text(x + 16, y + 18, "Notation", size=13, color=MUTED, bold=True)
    items = [
        ("PK", "Primary key", DOMAINS["payroll"], 110),
        ("FK", "Foreign key", DOMAINS["identity"], 110),
        ("UK", "Unique key", DOMAINS["hr"], 110),
        ("Crow’s foot", "many", LINE, 130),
        ("Circle", "optional (0..)", LINE, 140),
    ]
    ox = x + 16
    for badge, label, color, step in items:
        c.text(ox, y + 48, badge, size=13, color=color, bold=True)
        c.text(ox + 12 * len(badge) * 0.55 + 8, y + 48, label, size=13, color=INK)
        ox += step


def draw_table(c: Canvas, t: Table):
    color = DOMAINS[t.domain]
    c.rect(t.x, t.y, t.width, t.h, WHITE, STROKE, 1.4, radius=6)
    c.rect(t.x, t.y, t.width, HEADER_H, color, color, 1, radius=6)
    c.rect(t.x, t.y + HEADER_H - 8, t.width, 8, color, color, 0)
    c.text(t.x + 14, t.y + HEADER_H / 2, t.name, size=16, color=WHITE, bold=True)
    for i, col in enumerate(t.cols):
        yy = t.y + HEADER_H + i * ROW_H
        fill = WHITE
        if col.pk:
            fill = PK_TINT
        elif col.fk:
            fill = FK_TINT
        c.rect(t.x + 1, yy, t.width - 2, ROW_H, fill, fill, 0)
        badge = "PK" if col.pk else "FK" if col.fk else "UK" if col.uk else ""
        bcolor = DOMAINS["payroll"] if col.pk else DOMAINS["identity"] if col.fk else DOMAINS["hr"]
        if badge:
            c.text(t.x + 12, yy + ROW_H / 2, badge, size=12, color=bcolor, bold=True)
        c.text(t.x + (42 if badge else 12), yy + ROW_H / 2, col.name, size=13, color=INK, bold=col.pk)
        c.text(t.x + t.width - 12, yy + ROW_H / 2, col.typ, size=12, color=MUTED, anchor="end")
    c.rect(t.x, t.y, t.width, t.h, fill="none", stroke=STROKE, sw=1.4, radius=6)


def _anchor(t: Table, col: str, side: str):
    y = t.col_y(col)
    if side == "left":
        return t.x, y, -1, 0
    if side == "right":
        return t.right, y, 1, 0
    if side == "top":
        return t.cx, t.y, 0, -1
    return t.cx, t.bottom, 0, 1


def card_clearance(card: str) -> float:
    if card in ("0..1", "0..*"):
        return 26
    return 20


def draw_card(c: Canvas, x, y, nx, ny, card: str):
    size = 12
    tx, ty = -ny, nx
    ox, oy = x + nx * 1, y + ny * 1
    if card in ("0..*", "1..*", "*"):
        tip = (ox + nx * size, oy + ny * size)
        p0 = (ox + tx * size * 0.7, oy + ty * size * 0.7)
        p1 = (ox - tx * size * 0.7, oy - ty * size * 0.7)
        c.line([tip, p0], LINE, 1.8)
        c.line([tip, (ox, oy)], LINE, 1.8)
        c.line([tip, p1], LINE, 1.8)
        if card == "1..*":
            bx = ox + nx * (size + 5)
            by = oy + ny * (size + 5)
            c.line([(bx + tx * 7, by + ty * 7), (bx - tx * 7, by - ty * 7)], LINE, 2)
        elif card == "0..*":
            c.circle(ox + nx * (size + 7), oy + ny * (size + 7), 4, WHITE, LINE, 1.5)
    elif card in ("0..1", "1"):
        bx = ox + nx * 10
        by = oy + ny * 10
        c.line([(bx + tx * 8, by + ty * 8), (bx - tx * 8, by - ty * 8)], LINE, 2)
        if card == "0..1":
            c.circle(ox + nx * 20, oy + ny * 20, 4, WHITE, LINE, 1.5)
        else:
            bx2 = ox + nx * 16
            by2 = oy + ny * 16
            c.line([(bx2 + tx * 8, by2 + ty * 8), (bx2 - tx * 8, by2 - ty * 8)], LINE, 2)


def draw_rel(c: Canvas, tables: dict[str, Table], r: Rel):
    a, b = tables[r.a], tables[r.b]
    x1, y1, nx1, ny1 = _anchor(a, r.a_col, r.a_side)
    x2, y2, nx2, ny2 = _anchor(b, r.b_col, r.b_side)
    g1, g2 = card_clearance(r.a_card), card_clearance(r.b_card)
    p1 = (x1 + nx1 * g1, y1 + ny1 * g1)
    p2 = (x2 + nx2 * g2, y2 + ny2 * g2)
    if r.via:
        pts = [p1, *r.via, p2]
    elif r.a_side in ("left", "right") and r.b_side in ("left", "right"):
        mid = (p1[0] + p2[0]) / 2
        pts = [p1, (mid, p1[1]), (mid, p2[1]), p2]
    elif r.a_side in ("top", "bottom") and r.b_side in ("top", "bottom"):
        mid = (p1[1] + p2[1]) / 2
        pts = [p1, (p1[0], mid), (p2[0], mid), p2]
    else:
        if r.a_side in ("left", "right"):
            pts = [p1, (p2[0], p1[1]), p2]
        else:
            pts = [p1, (p1[0], p2[1]), p2]
    c.line(pts, LINE, 1.7)
    draw_card(c, x1, y1, nx1, ny1, r.a_card)
    draw_card(c, x2, y2, nx2, ny2, r.b_card)
    if r.label:
        mx = (p1[0] + p2[0]) / 2
        my = (p1[1] + p2[1]) / 2 - 10
        c.text(mx, my, r.label, size=12, color=MUTED, anchor="middle")


def tables_layout() -> list[Table]:
    """Spacious 5-column landscape layout with richer key columns."""
    T = lambda name, domain, cols, x, y, w=COL_W: Table(name, domain, cols, x, y, w)

    # Column X positions (identity | HR | payroll | academics | messaging)
    c1, c2, c3, c4, c5 = 50, 420, 820, 1220, 1680
    y0 = 150  # below domain banners

    return [
        # Identity
        T("roles", "identity", [PK(), C("name"), UK("slug")], c1, y0, 300),
        T(
            "users",
            "identity",
            [PK(), FK("role_id"), C("first_name"), C("last_name"), UK("email"), C("phone"), C("status")],
            c1,
            y0 + 160,
            300,
        ),
        T(
            "sessions",
            "identity",
            [PK("id", "varchar"), FK("user_id"), C("ip_address", "varchar(45)"), C("last_activity", "int")],
            c1,
            y0 + 420,
            300,
        ),
        T(
            "notifications",
            "governance",
            [PK(), FK("user_id"), C("type"), C("title"), C("severity"), C("is_read", "bool")],
            c1,
            y0 + 620,
            300,
        ),
        # HR
        T(
            "salary_grades",
            "hr",
            [PK(), UK("code"), C("name"), C("basic_salary", "decimal"), C("max_allowance_total", "decimal")],
            c2,
            y0,
            320,
        ),
        T(
            "staff",
            "hr",
            [
                PK(),
                FK("user_id"),
                FK("salary_grade_id"),
                UK("employee_id"),
                C("first_name"),
                C("last_name"),
                C("job_title"),
                C("department"),
                C("salary", "decimal"),
                C("ssnit_number"),
                C("bank_account"),
                C("status"),
            ],
            c2,
            y0 + 220,
            320,
        ),
        T(
            "allowance_types",
            "hr",
            [PK(), C("name"), UK("code"), C("is_taxable", "bool"), C("requires_authorization", "bool")],
            c2,
            y0 + 620,
            320,
        ),
        T(
            "staff_allowances",
            "hr",
            [PK(), FK("staff_id"), FK("allowance_type_id"), C("amount", "decimal"), C("is_authorized", "bool")],
            c2,
            y0 + 820,
            320,
        ),
        T(
            "loans",
            "hr",
            [
                PK(),
                FK("staff_id"),
                C("principal", "decimal"),
                C("outstanding_balance", "decimal"),
                C("monthly_deduction", "decimal"),
                C("status"),
            ],
            c2,
            y0 + 1020,
            320,
        ),
        # Payroll / leave / attendance
        T(
            "staff_attendances",
            "hr",
            [
                PK(),
                FK("staff_id"),
                UK("date", "date"),
                C("status"),
                C("penalty_amount", "decimal"),
                C("payroll_processed", "bool"),
            ],
            c3,
            y0,
            320,
        ),
        T(
            "leave_requests",
            "hr",
            [
                PK(),
                FK("staff_id"),
                C("leave_type"),
                C("start_date", "date"),
                C("end_date", "date"),
                C("status"),
                FK("reviewed_by"),
                FK("approved_by"),
            ],
            c3,
            y0 + 250,
            320,
        ),
        T(
            "payroll_runs",
            "payroll",
            [
                PK(),
                C("run_name"),
                C("pay_period_start", "date"),
                C("pay_period_end", "date"),
                C("payment_date", "date"),
                C("status"),
                C("total_net", "decimal"),
                FK("created_by"),
                FK("approved_by"),
            ],
            c3,
            y0 + 540,
            320,
        ),
        T(
            "payrolls",
            "payroll",
            [
                PK(),
                FK("payroll_run_id"),
                FK("staff_id"),
                C("basic_salary", "decimal"),
                C("allowances", "decimal"),
                C("income_tax", "decimal"),
                C("ssnit_contribution", "decimal"),
                C("loan_deductions", "decimal"),
                C("absence_penalties", "decimal"),
                C("net_salary", "decimal"),
                C("status"),
            ],
            c3,
            y0 + 860,
            320,
        ),
        T(
            "payroll_anomalies",
            "payroll",
            [
                PK(),
                C("scan_batch_id", "uuid"),
                FK("payroll_run_id"),
                FK("payroll_id"),
                FK("staff_id"),
                C("rule_code"),
                C("severity"),
                C("status"),
                FK("resolved_by"),
            ],
            c3,
            y0 + 1220,
            320,
        ),
        # Academics
        T(
            "school_classes",
            "academics",
            [PK(), C("name"), C("level"), C("sort_order", "smallint"), C("capacity", "int"), FK("teacher_id")],
            c4,
            y0,
            320,
        ),
        T(
            "students",
            "academics",
            [
                PK(),
                UK("admission_number"),
                C("first_name"),
                C("last_name"),
                C("gender"),
                FK("class_id"),
                FK("parent_id"),
                C("guardian_name"),
                C("status"),
            ],
            c4,
            y0 + 250,
            320,
        ),
        T(
            "subjects",
            "academics",
            [PK(), C("name"), UK("code", "varchar(20)"), C("levels", "json"), C("sort_order", "smallint")],
            c4,
            y0 + 560,
            320,
        ),
        T(
            "staff_subject",
            "academics",
            [PK(), FK("staff_id"), FK("subject_id")],
            c4,
            y0 + 780,
            320,
        ),
        T(
            "student_attendances",
            "academics",
            [PK(), FK("student_id"), FK("class_id"), UK("date", "date"), C("status"), FK("recorded_by")],
            c4,
            y0 + 940,
            320,
        ),
        T(
            "audit_logs",
            "governance",
            [PK(), FK("user_id"), C("action"), C("auditable_type"), C("auditable_id", "bigint"), C("ip_address", "varchar(45)")],
            c4,
            y0 + 1180,
            320,
        ),
        # Messaging + assessments
        T(
            "student_assessments",
            "academics",
            [
                PK(),
                FK("student_id"),
                FK("subject_id"),
                C("academic_year", "varchar(12)"),
                C("term", "tinyint"),
                C("classwork", "decimal"),
                C("project", "decimal"),
                C("assignment", "decimal"),
                C("homework", "decimal"),
                FK("recorded_by"),
            ],
            c5,
            y0,
            340,
        ),
        T(
            "student_term_reports",
            "academics",
            [
                PK(),
                FK("student_id"),
                C("academic_year", "varchar(12)"),
                C("term", "tinyint"),
                C("teacher_comment", "text"),
                FK("recorded_by"),
            ],
            c5,
            y0 + 360,
            340,
        ),
        T(
            "parent_messages",
            "messaging",
            [
                PK(),
                FK("sender_id"),
                C("type"),
                C("subject"),
                C("body", "text"),
                C("is_broadcast", "bool"),
                C("channels", "json"),
                C("sent_count", "int"),
            ],
            c5,
            y0 + 600,
            340,
        ),
        T(
            "parent_message_recipients",
            "messaging",
            [
                PK(),
                FK("parent_message_id"),
                FK("parent_id"),
                C("email"),
                C("phone"),
                C("status"),
                C("email_status"),
                C("whatsapp_status"),
                C("sent_at", "timestamp"),
            ],
            c5,
            y0 + 900,
            340,
        ),
    ]


def rels() -> list[Rel]:
    return [
        Rel("roles", "id", "users", "role_id", "1", "0..*", "bottom", "top"),
        Rel("users", "id", "staff", "user_id", "0..1", "0..1", "right", "left"),
        Rel("users", "id", "sessions", "user_id", "0..1", "0..*", "bottom", "top"),
        Rel("users", "id", "notifications", "user_id", "0..1", "0..*", "bottom", "top"),
        Rel("salary_grades", "id", "staff", "salary_grade_id", "0..1", "0..*", "bottom", "top"),
        Rel("staff", "id", "staff_allowances", "staff_id", "1", "0..*", "bottom", "top"),
        Rel("allowance_types", "id", "staff_allowances", "allowance_type_id", "1", "0..*", "bottom", "top"),
        Rel("staff", "id", "loans", "staff_id", "1", "0..*", "bottom", "top"),
        Rel("staff", "id", "staff_attendances", "staff_id", "1", "0..*", "right", "left"),
        Rel("staff", "id", "leave_requests", "staff_id", "1", "0..*", "right", "left"),
        Rel("staff", "id", "payrolls", "staff_id", "1", "0..*", "right", "left"),
        Rel("payroll_runs", "id", "payrolls", "payroll_run_id", "1", "1..*", "bottom", "top"),
        Rel("payrolls", "id", "payroll_anomalies", "payroll_id", "0..1", "0..*", "bottom", "top"),
        Rel("payroll_runs", "id", "payroll_anomalies", "payroll_run_id", "1", "0..*", "bottom", "top"),
        Rel("school_classes", "id", "students", "class_id", "0..1", "0..*", "bottom", "top"),
        Rel("staff", "id", "school_classes", "teacher_id", "0..1", "0..*", "right", "left", label="class tutor"),
        Rel("subjects", "id", "staff_subject", "subject_id", "1", "0..1", "bottom", "top"),
        Rel("staff", "id", "staff_subject", "staff_id", "1", "0..*", "right", "left"),
        Rel("students", "id", "student_attendances", "student_id", "1", "0..*", "bottom", "top"),
        Rel("students", "id", "student_assessments", "student_id", "1", "0..*", "right", "left"),
        Rel("students", "id", "student_term_reports", "student_id", "1", "0..*", "right", "left"),
        Rel("subjects", "id", "student_assessments", "subject_id", "1", "0..*", "right", "left"),
        Rel("parent_messages", "id", "parent_message_recipients", "parent_message_id", "1", "1..*", "bottom", "top"),
        Rel("users", "id", "students", "parent_id", "0..1", "0..*", "right", "left", label="parent"),
        Rel("users", "id", "parent_messages", "sender_id", "1", "0..*", "right", "left", label="sender"),
        Rel("users", "id", "audit_logs", "user_id", "0..1", "0..*", "right", "left"),
    ]


def build_diagram():
    # ~ A3 landscape at ~150 px/in design coords; exported at 300 DPI for Word
    width, height = 2140, 1760
    c = Canvas(
        width,
        height,
        "School SMS — Logical Data Model",
        "Entity-relationship overview of application tables (primary keys, foreign keys, uniqueness). "
        "Crow’s foot = many · circle = optional. Suitable for landscape Word / print.",
    )
    draw_title(c)

    # Domain banners
    banners = [
        (50, 105, 300, "Identity & alerts", "identity"),
        (420, 105, 320, "HR & compensation", "hr"),
        (820, 105, 320, "Attendance · leave · payroll", "payroll"),
        (1220, 105, 320, "Academics & audit", "academics"),
        (1680, 105, 340, "Assessment & messaging", "messaging"),
    ]
    for x, y, w, title, domain in banners:
        draw_domain_banner(c, x, y, w, title, domain)

    tables = tables_layout()
    lookup = {t.name: t for t in tables}
    for r in rels():
        draw_rel(c, lookup, r)
    for t in tables:
        draw_table(c, t)

    draw_legend(c, 50, height - 100)
    c.text(
        width - 40,
        height - 36,
        "Source: Laravel migrations · School SMS  ·  Insert in Word: Insert → Pictures → This Device (use landscape page)",
        size=12,
        color=MUTED,
        anchor="end",
    )
    c.text(
        50,
        height - 36,
        "Figure: Logical data model (enlarged for documentation)",
        size=12,
        color=MUTED,
    )
    return c


def write_word(png_path: Path):
    """Landscape Word document with the figure embedded at printable width."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Swap for landscape A4
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    run = title.add_run("School SMS — Logical Data Model")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    r = cap.add_run(
        "Entity-relationship overview of application tables. "
        "PK = primary key, FK = foreign key, UK = unique. Crow’s foot = many; circle = optional."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Usable width ≈ 29.7 - 3.0 = 26.7 cm
    doc.add_picture(str(png_path), width=Cm(26.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph()
    nr = note.add_run(
        "Note: Cross-domain foreign keys (e.g. parent_id, teacher_id, sender_id, recorded_by) "
        "are shown on subject-area diagrams where space allows. "
        "Source: backend/database/migrations."
    )
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    docx_path = OUT / "02-logical-overview-word.docx"
    doc.save(docx_path)
    return docx_path


def patch_index():
    index = OUT / "index.html"
    if not index.exists():
        return
    html = index.read_text(encoding="utf-8")
    card = """
            <article class="card" style="border-color:#1B365D;">
              <h2>Logical data model (Word / print)</h2>
              <p>Enlarged ER diagram with domain bands, richer columns, 300&nbsp;DPI PNG and a ready landscape Word file.</p>
              <div class="actions">
                <a href="02-logical-overview-word.docx" download>Download Word (.docx)</a>
                <a href="02-logical-overview-word.png" download>Download PNG (300 DPI)</a>
                <a href="02-logical-overview-word.svg" download>Download SVG</a>
              </div>
              <a href="02-logical-overview-word.svg" target="_blank"><img src="02-logical-overview-word.svg" alt="Logical data model Word"></a>
            </article>"""
    if "02-logical-overview-word" in html:
        return
    # Insert after the existing Logical overview card if present
    marker = "<h2>Logical overview</h2>"
    pos = html.find(marker)
    if pos != -1:
        # find end of that article
        end = html.find("</article>", pos)
        if end != -1:
            html = html[: end + 10] + card + html[end + 10 :]
            index.write_text(html, encoding="utf-8")
            return
    end = html.find("</p>", html.find('<p class="lead">'))
    if end != -1:
        index.write_text(html[: end + 4] + card + html[end + 4 :], encoding="utf-8")


def main():
    c = build_diagram()
    # Word-specific stems
    svg, png = c.save("02-logical-overview-word")
    # Also replace the standard logical overview with the enlarged master
    c.save("02-logical-overview")
    docx = write_word(png)
    patch_index()
    print("Wrote:")
    print(" ", png, f"({png.stat().st_size // 1024} KB, {DPI} DPI)")
    print(" ", svg)
    print(" ", docx)
    print("  02-logical-overview.png/svg (enlarged master)")


if __name__ == "__main__":
    main()
